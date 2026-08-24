# -*- coding: utf-8 -*-
"""Снять оформление с DOCX-образца: отчет для человека и профиль для сборки.

    python inspect_sample.py sample.docx                  - отчет в консоль
    python inspect_sample.py sample.docx --json p.json     - плюс профиль сборки
    python inspect_sample.py sample.docx --paragraphs 80   - показать больше абзацев

Отчет отвечает на вопросы: какие стили реально используются и на каком тексте,
как устроены секции и колонтитулы, чем отличаются таблицы (стиль, ширины, заливка
шапки, границы, повтор шапки, стили абзацев в ячейках).
"""
import argparse
import json
import os
import re
import sys
import zipfile

from docx import Document
from docx.enum.section import WD_ORIENT

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def cm(v):
    return round(v.cm, 2) if v is not None else None


def para_num(p):
    numPr = p._p.find(".//" + W + "numPr")
    if numPr is None:
        return None
    numId = numPr.find(W + "numId")
    return numId.get(W + "val") if numId is not None else "?"


def used_styles(doc):
    """Стиль -> сколько раз, где (текст или таблица) и примеры текста.

    Абзацы внутри таблиц в doc.paragraphs не попадают, а именно там живут стили
    шапки и тела таблиц - их надо собирать отдельно.
    """
    out = {}

    def add(p, where, head=False):
        if not p.text.strip():
            return
        name = p.style.name if p.style is not None else "(нет стиля)"
        rec = out.setdefault(name, {"count": 0, "samples": [], "unnumbered": 0,
                                    "where": set(), "table_head": 0})
        rec["count"] += 1
        rec["where"].add(where)
        if head:
            rec["table_head"] += 1
        if para_num(p) == "0":
            rec["unnumbered"] += 1
        if len(rec["samples"]) < 4:
            rec["samples"].append(p.text.strip()[:58])

    for p in doc.paragraphs:
        add(p, "текст")
    for t in doc.tables:
        for ri, row in enumerate(t.rows):
            for c in row.cells:
                for p in c.paragraphs:
                    add(p, "таблица", head=(ri == 0))
    return out


def style_params(path):
    """Параметры пользовательских стилей прямо из styles.xml."""
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/styles.xml").decode("utf-8")
    out = {}
    for m in re.finditer(r'<w:style [^>]*w:styleId="([^"]+)"[^>]*>(.*?)</w:style>', xml, re.S):
        sid, body = m.group(1), m.group(2)
        nm = re.search(r'<w:name w:val="([^"]+)"', body)
        if not nm:
            continue
        info = {"id": sid}
        f = re.search(r'<w:rFonts[^>]*w:ascii="([^"]+)"', body)
        if f:
            info["font"] = f.group(1)
        s = re.search(r'<w:sz w:val="(\d+)"', body)
        if s:
            info["size_pt"] = int(s.group(1)) / 2.0
        if "<w:b/>" in body:
            info["bold"] = True
        sp = re.search(r'<w:spacing([^>]*)/>', body)
        if sp:
            info["spacing"] = sp.group(1).strip()
        ind = re.search(r'<w:ind([^>]*)/>', body)
        if ind:
            info["indent"] = ind.group(1).strip()
        jc = re.search(r'<w:jc w:val="(\w+)"', body)
        if jc:
            info["align"] = jc.group(1)
        num = re.search(r'<w:numId w:val="(\d+)"', body)
        if num:
            info["numId"] = num.group(1)
        lvl = re.search(r'<w:outlineLvl w:val="(\d+)"', body)
        if lvl:
            info["outline"] = int(lvl.group(1))
        based = re.search(r'<w:basedOn w:val="([^"]+)"', body)
        if based:
            info["based_on_id"] = based.group(1)
        out[nm.group(1)] = info
    # Уровень структуры наследуется от базового стиля: без этого заголовок
    # третьего уровня выглядит как обычный текст.
    by_id = {v["id"]: k for k, v in out.items()}
    for name, info in out.items():
        seen = 0
        cur = info
        while "outline" not in cur and cur.get("based_on_id") and seen < 6:
            parent = out.get(by_id.get(cur["based_on_id"], ""), None)
            if not parent:
                break
            if "outline" in parent:
                info["outline_inherited"] = parent["outline"]
                break
            cur = parent
            seen += 1
    return out


def table_info(doc):
    rows = []
    for i, t in enumerate(doc.tables, 1):
        pr = t._tbl.find(W + "tblPr")
        st = pr.find(W + "tblStyle") if pr is not None else None
        layout = pr.find(W + "tblLayout") if pr is not None else None
        borders = pr.find(W + "tblBorders") if pr is not None else None
        head_cell = t.rows[0].cells[0]
        shd = re.search(r'<w:shd[^>]*w:fill="([0-9A-Fa-f]{6}|auto)"', head_cell._tc.xml)
        widths = [cm(c.width) for c in t.rows[0].cells]
        cell_styles = []
        for ri in range(min(2, len(t.rows))):
            names = []
            for c in t.rows[ri].cells[:3]:
                for p in c.paragraphs[:1]:
                    names.append(p.style.name if p.style is not None else "-")
            cell_styles.append("/".join(names))
        border_sz = None
        if borders is not None:
            top = borders.find(W + "top")
            border_sz = top.get(W + "sz") if top is not None else None
        full_width = []
        for r in t.rows[1:]:
            texts = [c.text.strip() for c in r.cells]
            if texts and len(set(texts)) == 1 and texts[0]:
                full_width.append(texts[0])
        rows.append({
            "index": i,
            "full_width_rows": full_width,
            "style_id": st.get(W + "val") if st is not None else None,
            "style_name": (t.style.name if t.style is not None else None),
            "cols": len(t.columns),
            "rows": len(t.rows),
            "layout": layout.get(W + "type") if layout is not None else None,
            "head_fill": shd.group(1) if shd else None,
            "border_sz": border_sz,
            "repeat_header": t.rows[0]._tr.find(".//" + W + "tblHeader") is not None,
            "widths_cm": widths,
            "head_text": " | ".join(c.text.strip()[:18] for c in t.rows[0].cells[:4]),
            "cell_styles": cell_styles,
        })
    return rows


def hf_texts(path):
    """Текст колонтитулов прямо из файла.

    Через python-docx он часто не виден: в шаблонах колонтитул обернут в sdt,
    и section.footer.paragraphs оказывается пустым.
    """
    out = {}
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if not re.match(r"word/(header|footer)\d*\.xml", name):
                continue
            xml = z.read(name).decode("utf-8")
            parts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
            txt = " ".join(p for p in parts if p.strip())
            if txt.strip():
                out[os.path.basename(name)] = txt.strip()
    return out


def sections_info(doc):
    out = []
    for i, s in enumerate(doc.sections, 1):
        foot = " / ".join(p.text.strip() for p in s.footer.paragraphs if p.text.strip())
        head = " / ".join(p.text.strip() for p in s.header.paragraphs if p.text.strip())
        out.append({
            "index": i,
            "orientation": "landscape" if s.orientation == WD_ORIENT.LANDSCAPE else "portrait",
            "page_cm": [cm(s.page_width), cm(s.page_height)],
            "margins_cm": [cm(s.left_margin), cm(s.right_margin), cm(s.top_margin), cm(s.bottom_margin)],
            "first_page_differs": s.different_first_page_header_footer,
            "footer": foot,
            "header": head,
        })
    return out


def guess_profile(path, styles, tables, sections, params, hf=None):
    """Черновой профиль сборки. Требует ручной доводки под конкретный жанр."""
    def outline_of(name):
        info = params.get(name, {})
        if "outline" in info:
            return info["outline"]
        return info.get("outline_inherited")

    def head_style(level):
        cands = [n for n in styles if outline_of(n) == level]
        if not cands:
            return None
        return max(cands, key=lambda n: styles[n]["count"])

    # Стиль по умолчанию берем только если своих стилей у образца нет: в таблицах
    # он лидирует по частоте за счет титульных блоков и забивает настоящие стили.
    DEFAULTS = {"normal", "обычный", "default paragraph font"}

    def rank(cands):
        own = [c for c in cands if c[1].lower() not in DEFAULTS]
        pool = own or cands
        pool.sort(reverse=True)
        return pool[0][1] if pool else None

    def table_style(head):
        """Стиль абзацев в шапке таблиц либо в их теле."""
        cands = []
        for name, rec in styles.items():
            if "таблица" not in rec["where"]:
                continue
            in_head = rec["table_head"]
            if head and in_head:
                cands.append((in_head, name))
            elif not head and rec["count"] > in_head:
                cands.append((rec["count"] - in_head, name))
        return rank(cands)

    def body_style():
        cands = [(rec["count"], n) for n, rec in styles.items()
                 if "текст" in rec["where"] and outline_of(n) is None
                 and not params.get(n, {}).get("numId")]
        return rank(cands)

    def bullet_style():
        for name, info in params.items():
            if info.get("numId") and outline_of(name) is None and "список" in name.lower():
                return name
        return None

    # второй по частоте стиль тела таблиц: часто отличается кеглем
    table_bodies = [n for n, rec in styles.items()
                    if "таблица" in rec["where"] and rec["count"] > rec["table_head"]
                    and n.lower() not in DEFAULTS]
    table_bodies.sort(key=lambda n: -styles[n]["count"])

    body = body_style()
    heads = [head_style(0), head_style(1), head_style(2)]
    prof = {
        "sample": path,
        "page": {
            "width_cm": sections[0]["page_cm"][0] if sections else 21.0,
            "height_cm": sections[0]["page_cm"][1] if sections else 29.7,
            "margins_cm": sections[0]["margins_cm"] if sections else [2.5, 2.0, 2.0, 2.0],
        },
        "styles": {
            "body": body,
            "h1": heads[0],
            "h2": heads[1],
            "h3": heads[2],
            "callout": next((n for n, rec in styles.items()
                             if "текст" in rec["where"] and outline_of(n) is None
                             and params.get(n, {}).get("bold") and n != body), None),
            "bullet": bullet_style(),
            "table_head": table_style(head=True),
            "table_body": table_bodies[0] if table_bodies else None,
            "table_body_alt": table_bodies[1] if len(table_bodies) > 1 else None,
        },
        "tables": {
            "default_style_id": None,
            "grid_style_id": None,
            "head_fill": None,
            "head_fill_alt": None,
            "border_sz": 4,
            "repeat_header": True,
            "numbering_dot": True,
            "counter_headers": ["№"],
            "counter_widths": {"default": 1.1, "long_header": 1.4},
        },
        "footer": {"line": True, "size_pt": 8, "line_len": 77},
        "title_page": {"separate_section": True, "no_footer": True},
        "toc": {"placeholder": "Для сборки оглавления выделить документ целиком и нажать F9."},
        "headings_map": {"h1": [], "h1_unnumbered": [], "h2": [], "h3": [], "callout": []},
    }
    fills = [t["head_fill"] for t in tables if t["head_fill"] and t["head_fill"] != "auto"]
    if fills:
        prof["tables"]["head_fill"] = max(set(fills), key=fills.count)
        others = [f for f in fills if f != prof["tables"]["head_fill"]]
        if others:
            prof["tables"]["head_fill_alt"] = max(set(others), key=others.count)
    ids = [t["style_id"] for t in tables if t["style_id"]]
    if ids:
        prof["tables"]["default_style_id"] = max(set(ids), key=ids.count)
        rest = [i for i in ids if i != prof["tables"]["default_style_id"]]
        if rest:
            prof["tables"]["grid_style_id"] = max(set(rest), key=rest.count)
    szs = [int(t["border_sz"]) for t in tables if t["border_sz"]]
    if szs:
        prof["tables"]["border_sz"] = max(set(szs), key=szs.count)
    if any(t["full_width_rows"] for t in tables):
        marks = sorted({m for t in tables for m in t["full_width_rows"]})
        prof["tables"]["full_width_rows"] = marks
    foot = next((s["footer"] for s in sections if s["footer"]), "")
    if not foot and hf:
        foot = next((v for k, v in sorted(hf.items()) if k.startswith("footer")), "")
    if foot:
        # В подвале образца стоит название чужого документа: заменяем его подстановкой.
        prof["footer"]["template"] = re.sub(r"\s{2,}", " ", re.sub(r"«[^»]*»", "«{title}»", foot)).strip("_ ")
    return prof


def main():
    # Вывод содержит кириллицу. Без явного переключения печать падает с UnicodeEncodeError
    # везде, где консоль не в UTF-8: сборочный агент, чужая локаль.
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
    ap = argparse.ArgumentParser()
    ap.add_argument("sample")
    ap.add_argument("--json", help="куда записать черновой профиль")
    ap.add_argument("--paragraphs", type=int, default=40, help="сколько абзацев показать")
    a = ap.parse_args()

    doc = Document(a.sample)
    params = style_params(a.sample)
    styles = used_styles(doc)
    tables = table_info(doc)
    sections = sections_info(doc)
    hf = hf_texts(a.sample)

    print("=== СЕКЦИИ ===")
    for s in sections:
        print("  %d) %s %sx%s см, поля L%s R%s T%s B%s, первая страница отдельно: %s"
              % (s["index"], s["orientation"], s["page_cm"][0], s["page_cm"][1],
                 s["margins_cm"][0], s["margins_cm"][1], s["margins_cm"][2], s["margins_cm"][3],
                 s["first_page_differs"]))
        if s["header"]:
            print("       верхний колонтитул: %s" % s["header"][:70])
        if s["footer"]:
            print("       нижний колонтитул: %s" % s["footer"][:70])

    print()
    print("=== СТИЛИ АБЗАЦЕВ В ДЕЛЕ ===")
    for name, rec in sorted(styles.items(), key=lambda kv: -kv[1]["count"]):
        info = params.get(name, {})
        bits = ["в " + "+".join(sorted(rec["where"]))]
        for k in ("font", "size_pt", "bold", "align", "numId", "outline", "outline_inherited"):
            if k in info:
                bits.append("%s=%s" % (k, info[k]))
        if rec["table_head"]:
            bits.append("шапка таблиц x%d" % rec["table_head"])
        if rec["unnumbered"]:
            bits.append("без номера x%d" % rec["unnumbered"])
        print("  %-30s x%-3d %s" % (name[:30], rec["count"], ", ".join(bits)))
        for s in rec["samples"]:
            print("        %s" % s)

    print()
    print("=== ТАБЛИЦЫ ===")
    for t in tables:
        print("  %2d) %dx%d стиль=%s заливка шапки=%s границы sz=%s повтор шапки=%s"
              % (t["index"], t["rows"], t["cols"], t["style_id"], t["head_fill"],
                 t["border_sz"], t["repeat_header"]))
        print("       ширины: %s" % t["widths_cm"])
        print("       шапка: %s" % t["head_text"])
        print("       стили ячеек: %s" % " ; ".join(t["cell_styles"]))

    print()
    print("=== ПОРЯДОК БЛОКОВ (первые %d) ===" % a.paragraphs)
    body = doc.element.body
    shown = 0
    ti = 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "tbl":
            ti += 1
            if ti <= len(tables):
                print("  [таблица %d] %dx%d %s" % (ti, tables[ti - 1]["rows"],
                                                   tables[ti - 1]["cols"],
                                                   tables[ti - 1]["head_text"][:46]))
            continue
        if tag != "p" or shown >= a.paragraphs:
            continue
        for p in doc.paragraphs:
            if p._p is child:
                if p.text.strip():
                    shown += 1
                    print("  %-26s %s" % ((p.style.name if p.style else "-")[:26],
                                          p.text.strip()[:74]))
                break

    if a.json:
        prof = guess_profile(a.sample, styles, tables, sections, params, hf)
        with open(a.json, "w", encoding="utf-8") as f:
            json.dump(prof, f, ensure_ascii=False, indent=2)
        print()
        print("Черновой профиль записан: %s" % a.json)
        print("Проверить руками по отчету выше: styles.h1/h2/h3, что считать основным стилем "
              "таблиц (table_body против table_body_alt), tables.default_style_id и "
              "grid_style_id, headings_map, footer.template.")


if __name__ == "__main__":
    sys.exit(main())
