# -*- coding: utf-8 -*-
"""Сборка DOCX по образцу: образец служит шаблоном стилей, содержимое подставляется новое.

Образец открывается как документ, его тело очищается, стили / нумерация / темы / параметры
страницы остаются родными. Дальше документ наполняется через методы SampleDoc.

Использование:

    from docx_builder import SampleDoc
    d = SampleDoc("profile.json")
    d.title_block([("Отчет", 22, True), ("Проект N", 16, False)])
    d.section("portrait")
    d.heading("Введение", 1)
    d.body("Текст абзаца.")
    d.table(["№ п/п", "Показатель", "Значение"], rows)
    d.save("out.docx", author="Имя Ф.", title="Отчет")
"""
import copy
import json
import os

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

EMU_PER_CM = 360000


class SampleDoc:
    """Документ, собираемый по образцу."""

    def __init__(self, profile, sample=None):
        profile_dir = None
        if isinstance(profile, str):
            profile_dir = os.path.dirname(os.path.abspath(profile))
            with open(profile, encoding="utf-8") as f:
                profile = json.load(f)
        self.p = profile
        self.sample = self._resolve_sample(sample or profile.get("sample"), profile_dir)
        self.doc = Document(self.sample)
        self.warnings = []
        page = self.p.get("page", {})
        self.pw = Cm(page.get("width_cm", 21.0))
        self.ph = Cm(page.get("height_cm", 29.7))
        m = page.get("margins_cm", [2.5, 2.0, 2.0, 2.0])
        self.m_left, self.m_right, self.m_top, self.m_bottom = [Cm(x) for x in m]
        self.band_portrait = page.get("width_cm", 21.0) - m[0] - m[1]
        self.band_landscape = page.get("height_cm", 29.7) - m[0] - m[1]
        self.band = self.band_portrait
        self.footer_text = None
        self._clear_body()

    # ------------------------------------------------------------- внутреннее

    @staticmethod
    def _resolve_sample(path, profile_dir):
        """Образец задается профилем, но путь может быть и относительным.

        Профиль переезжает между машинами и папками вместе с образцом, поэтому
        относительный путь ищется рядом с самим профилем.
        """
        if not path:
            raise SystemExit("в профиле не задан образец (ключ sample) и не передан явно")
        if os.path.isabs(path):
            if os.path.exists(path):
                return path
            raise SystemExit("образец не найден: %s" % path)
        for base in ([profile_dir] if profile_dir else []) + [os.getcwd()]:
            cand = os.path.join(base, path)
            if os.path.exists(cand):
                return cand
        raise SystemExit("образец не найден ни рядом с профилем, ни в текущей папке: %s" % path)

    def _clear_body(self):
        body = self.doc.element.body
        sect = copy.deepcopy(body.find(qn("w:sectPr")))
        for el in list(body):
            body.remove(el)
        body.append(sect)
        sec = self.doc.sections[0]
        self._apply_page(sec, landscape=False)
        if self.p.get("title_page", {}).get("no_footer", True):
            sec.different_first_page_header_footer = True

    def _apply_page(self, sec, landscape):
        if landscape:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = self.ph, self.pw
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width, sec.page_height = self.pw, self.ph
        sec.left_margin, sec.right_margin = self.m_left, self.m_right
        sec.top_margin, sec.bottom_margin = self.m_top, self.m_bottom

    def _style(self, key, required=False):
        """Имя стиля из профиля; если стиля нет в образце - предупреждение и None."""
        name = self.p.get("styles", {}).get(key)
        if not name:
            if required:
                self.warnings.append("в профиле нет стиля %r" % key)
            return None
        try:
            self.doc.styles[name]
        except KeyError:
            self.warnings.append("стиля %r нет в образце (ключ %s)" % (name, key))
            return None
        return name

    def _style_by_id(self, style_id):
        if not style_id:
            return None
        for s in self.doc.styles:
            if s.style_id == style_id:
                return s
        self.warnings.append("стиля таблицы с id %r нет в образце" % style_id)
        return None

    # ------------------------------------------------------------- абзацы

    def para(self, text="", style_key=None, style_name=None, align=None, size=None,
             bold=None, keep_next=False, space_after=None):
        name = style_name if style_name else (self._style(style_key) if style_key else None)
        p = self.doc.add_paragraph(style=name)
        if keep_next:
            p.paragraph_format.keep_with_next = True
        if text:
            r = p.add_run(text)
            if size:
                r.font.size = Pt(size)
            if bold is not None:
                r.bold = bold
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    def body(self, text):
        return self.para(text, style_key="body")

    def callout(self, text, keep_next=True):
        return self.para(text, style_key="callout", keep_next=keep_next)

    def bullet(self, text):
        return self.para(text, style_key="bullet")

    def heading(self, text, level=1, numbered=True, keep_next=False):
        p = self.para(text, style_key="h%d" % level, keep_next=keep_next)
        if not numbered:
            self.unnumber(p)
        return p

    @staticmethod
    def unnumber(p):
        """Снять нумерацию с абзаца, стиль которого нумерованный (numId=0)."""
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "0")
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        return p

    def title_block(self, lines):
        """Титульные строки: список (текст, кегль, жирный)."""
        for text, size, bold in lines:
            self.para(text, align="center", size=size, bold=bold)

    def empty(self, count=1):
        for _ in range(count):
            self.para()

    def toc(self, levels="1-3", placeholder=None):
        """Поле оглавления. Word собирает его по F9; до этого виден placeholder."""
        placeholder = placeholder or self.p.get("toc", {}).get(
            "placeholder", "Для сборки оглавления выделить документ целиком и нажать F9.")
        p = self.doc.add_paragraph()
        r = p.add_run()
        beg = OxmlElement("w:fldChar")
        beg.set(qn("w:fldCharType"), "begin")
        r._r.append(beg)
        r2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "%s" \\h \\z \\u ' % levels
        r2._r.append(instr)
        r3 = p.add_run()
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r3._r.append(sep)
        p.add_run(placeholder)
        r5 = p.add_run()
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r5._r.append(end)
        return p

    # ------------------------------------------------------------- секции

    def section(self, orientation="portrait", footer=True, footer_text=None):
        """Новая секция с новой страницы. Ориентация меняет полосу набора."""
        sec = self.doc.add_section(WD_SECTION.NEW_PAGE)
        # Свойство наследуется от предыдущей секции: без сброса подвал пропадет
        # на первой странице КАЖДОЙ секции.
        sec.different_first_page_header_footer = False
        landscape = orientation == "landscape"
        self._apply_page(sec, landscape)
        self.band = self.band_landscape if landscape else self.band_portrait
        if footer:
            self.set_footer(sec, footer_text or self.footer_text)
        return sec

    def set_footer(self, sec, text):
        if not text:
            return
        cfg = self.p.get("footer", {})
        size = cfg.get("size_pt", 8)
        sec.footer.is_linked_to_previous = False
        f = sec.footer
        for p in list(f.paragraphs)[1:]:
            p._p.getparent().remove(p._p)
        first = f.paragraphs[0]
        first.text = ""
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first.paragraph_format.space_after = Pt(0)
        if cfg.get("line", True):
            first.add_run("_" * cfg.get("line_len", 77)).font.size = Pt(size)
            second = f.add_paragraph()
        else:
            second = first
        second.alignment = WD_ALIGN_PARAGRAPH.CENTER
        second.paragraph_format.space_after = Pt(0)
        second.add_run(text).font.size = Pt(size)

    def footer_for_all(self, text):
        """Текст подвала для секций, создаваемых дальше."""
        self.footer_text = text

    # ------------------------------------------------------------- таблицы

    def counter_width(self, header):
        cfg = self.p.get("tables", {}).get("counter_widths", {})
        if header in cfg:
            return cfg[header]
        if len(header) > 5:
            return cfg.get("long_header", 1.4)
        return cfg.get("default", 1.1)

    def is_counter(self, header, values):
        marks = self.p.get("tables", {}).get("counter_headers", ["№"])
        if any(header.startswith(m) for m in marks):
            return True
        if header in self.p.get("tables", {}).get("counter_widths", {}):
            return True
        if len(header) > 14:
            return False
        seen = 0
        for v in values:
            v = str(v).strip()
            if not v:
                continue
            if not v.rstrip(".").isdigit() or len(v) > 4:
                return False
            seen += 1
        return seen > 0

    def auto_widths(self, header, rows, band=None, min_w=1.6):
        """Ширины по содержимому: колонка-счетчик узкая, остальные делят остаток.

        Вес колонки - корень из длины самой длинной ячейки: без корня одна
        многострочная ячейка забирает всю таблицу.
        """
        band = band or self.band
        n = len(header)
        cols = [(header[i], [str(r[i]) if i < len(r) else "" for r in rows]) for i in range(n)]
        widths = [None] * n
        free = band
        flex = []
        for i, (h, vals) in enumerate(cols):
            if self.is_counter(h, vals):
                widths[i] = self.counter_width(h)
                free -= widths[i]
            else:
                flex.append(i)
        weights = []
        for i in flex:
            h, vals = cols[i]
            longest = max([len(h)] + [len(v) for v in vals]) if vals else len(h)
            weights.append(min(longest, 400) ** 0.5)
        total = sum(weights) or 1
        for k, i in enumerate(flex):
            widths[i] = max(min_w, free * weights[k] / total)
        over = sum(widths) - band
        if over > 0.01 and flex:
            widest = max(flex, key=lambda i: widths[i])
            widths[widest] -= over
        return widths

    def table(self, header, rows, widths=None, style="default", head_fill=None,
              body_style_key="table_body", head_style_key="table_head", borders=None,
              repeat=None, head_size=None, merge_rows=None, numbering_dot=None):
        """Таблица в оформлении образца.

        style: "default" | "grid" | None (без табличного стиля)
        merge_rows: список (индекс строки данных, текст) - строка-разделитель на всю ширину
        """
        tbl_cfg = self.p.get("tables", {})
        rows = [list(r) for r in rows]
        if numbering_dot is None:
            numbering_dot = tbl_cfg.get("numbering_dot", True)
        if numbering_dot and header and str(header[0]).startswith("№"):
            for r in rows:
                if r and str(r[0]).isdigit():
                    r[0] = "%s." % r[0]
        t = self.doc.add_table(rows=0, cols=len(header))
        style_id = None
        if style == "default":
            style_id = tbl_cfg.get("default_style_id")
        elif style == "grid":
            style_id = tbl_cfg.get("grid_style_id")
        st = self._style_by_id(style_id) if style_id else None
        if st is not None:
            t.style = st
        if widths is None:
            widths = self.auto_widths(header, rows)
        fill = head_fill if head_fill is not None else tbl_cfg.get("head_fill", "D9D9D9")
        head_name = self._style(head_style_key)
        body_name = self._style(body_style_key)

        tr = t.add_row()
        for i, h in enumerate(header):
            cell = tr.cells[i]
            if fill:
                self.shade(cell, fill)
            p = cell.paragraphs[0]
            if head_name:
                p.style = self.doc.styles[head_name]
            r = p.add_run(str(h))
            if head_size:
                r.font.size = Pt(head_size)
                r.bold = True
        for row in rows:
            tr = t.add_row()
            for i, val in enumerate(row):
                if i >= len(header):
                    break
                cell = tr.cells[i]
                self.shade(cell, "auto")
                p = cell.paragraphs[0]
                if body_name:
                    p.style = self.doc.styles[body_name]
                if val:
                    p.add_run(str(val))
        self.fixed_layout(t)
        self.set_widths(t, widths)
        sz = borders if borders is not None else tbl_cfg.get("border_sz", 4)
        if sz:
            self.set_borders(t, sz)
        if repeat if repeat is not None else tbl_cfg.get("repeat_header", True):
            self.repeat_header(t)
        for idx, label in (merge_rows or []):
            self.merge_full_row(t, idx + 1, label, body_name)
        # Две таблицы подряд Word склеивает в одну - разделяем пустым абзацем.
        self.para()
        return t

    def merge_full_row(self, table, row_index, label, style_name=None):
        row = table.rows[row_index]
        cell = row.cells[0].merge(row.cells[-1])
        cell.text = ""
        p = cell.paragraphs[0]
        if style_name:
            p.style = self.doc.styles[style_name]
        p.add_run(label).bold = True
        return cell

    @staticmethod
    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:color"), "auto")
        el.set(qn("w:fill"), fill)
        tcPr.append(el)

    @staticmethod
    def set_borders(table, sz=4):
        tblPr = table._tbl.tblPr
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        b = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(sz))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "auto")
            b.append(e)
        tblPr.append(b)

    @staticmethod
    def fixed_layout(table):
        tblPr = table._tbl.tblPr
        for old in tblPr.findall(qn("w:tblLayout")):
            tblPr.remove(old)
        el = OxmlElement("w:tblLayout")
        el.set(qn("w:type"), "fixed")
        tblPr.append(el)

    @staticmethod
    def repeat_header(table):
        trPr = table.rows[0]._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:tblHeader"))

    @staticmethod
    def set_widths(table, widths_cm):
        table.autofit = False
        for i, w in enumerate(widths_cm):
            if i < len(table.columns):
                table.columns[i].width = Cm(w)
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths_cm):
                    cell.width = Cm(widths_cm[i])

    def rule_table(self, cols=3, border_sz=24):
        """Декоративная таблица-линейка (часто стоит на титуле образцов)."""
        t = self.doc.add_table(rows=1, cols=cols)
        self.set_widths(t, [self.band / float(cols)] * cols)
        self.fixed_layout(t)
        b = OxmlElement("w:tblBorders")
        e = OxmlElement("w:bottom")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(border_sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        b.append(e)
        t._tbl.tblPr.append(b)
        self.para()
        return t

    def plain_table(self, rows, widths=None, bold_first_row=False):
        """Таблица без стиля и без границ (шапка документа, блок подписей)."""
        if not rows:
            return None
        t = self.doc.add_table(rows=0, cols=len(rows[0]))
        for ri, row in enumerate(rows):
            tr = t.add_row()
            for i, val in enumerate(row):
                p = tr.cells[i].paragraphs[0]
                r = p.add_run(str(val))
                if bold_first_row and ri == 0:
                    r.bold = True
        self.set_widths(t, widths or [self.band / float(len(rows[0]))] * len(rows[0]))
        self.para()
        return t

    # ------------------------------------------------------------- сохранение

    def save(self, path, author=None, title=None):
        if author:
            self.doc.core_properties.author = author
            self.doc.core_properties.last_modified_by = author
        if title:
            self.doc.core_properties.title = title
        self.doc.save(path)
        return path
