# -*- coding: utf-8 -*-
"""Проверка собранного DOCX: содержимое, оформление, поля Word, картинки страниц.

    python verify_result.py out.docx --md out.md            - сверка с исходником
    python verify_result.py out.docx --sample sample.docx   - сравнить оформление с образцом
    python verify_result.py out.docx --update-fields        - собрать оглавление (нужен Word)
    python verify_result.py out.docx --png 1,2,9            - картинки страниц для глаз

Сверка с markdown ловит потерю строк и склейку таблиц: число таблиц и строк в DOCX
должно совпадать с числом блоков и строк в исходнике.
Картинки страниц - единственный способ увидеть реальную верстку: разрывы, вылезание
таблиц за поля, пропавшие колонтитулы.
"""
import argparse
import glob
import io
import os
import re
import subprocess
import sys

from docx import Document
from docx.enum.section import WD_ORIENT

SEP = re.compile(r"^\|[\s\-:|]*-[\s\-:|]*\|$")
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
# Символы заданы кодами, чтобы сам файл детектора их не содержал.
BAD = {"\u2014": "тире длинное", "\u2013": "тире короткое",
       "\u0451": "е с точками", "\u0401": "Е с точками",
       "\u2026": "многоточие", "\u2192": "стрелка"}

PS_UPDATE = r'''
$OutputEncoding = [Console]::OutputEncoding = [Text.Encoding]::UTF8
$ErrorActionPreference = "Stop"
$was = Get-Process winword -ErrorAction SilentlyContinue
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open("{docx}", $false, $false)
$doc.Fields.Update() | Out-Null
if ($doc.TablesOfContents.Count -gt 0) {{ $doc.TablesOfContents.Item(1).Update() | Out-Null }}
$doc.Repaginate()
"страниц: $($doc.ComputeStatistics(2))"
$doc.Save()
{export}
$doc.Close($false)
if (-not $was) {{ $word.Quit() }}
'''


def md_stats(md_path):
    tables = rows = 0
    in_tab = False
    with io.open(md_path, encoding="utf-8") as f:
        for ln in f.read().split("\n"):
            if ln.startswith("|"):
                if not in_tab:
                    tables += 1
                    in_tab = True
                if not SEP.match(ln.strip()):
                    rows += 1
            else:
                in_tab = False
    return tables, rows


def docx_report(path):
    d = Document(path)
    tables = [(len(t.columns), len(t.rows)) for t in d.tables]
    rows = sum(r for _, r in tables)
    styles = {}
    for p in d.paragraphs:
        if p.text.strip():
            nm = p.style.name if p.style is not None else "(нет стиля)"
            styles[nm] = styles.get(nm, 0) + 1
    secs = []
    for s in d.sections:
        band = (s.page_width.cm - s.left_margin.cm - s.right_margin.cm)
        foot = any(p.text.strip() for p in s.footer.paragraphs)
        secs.append({"orient": "альб" if s.orientation == WD_ORIENT.LANDSCAPE else "кн",
                     "band": round(band, 1), "footer": foot,
                     "first_differs": s.different_first_page_header_footer})
    over = []
    for i, t in enumerate(d.tables, 1):
        wsum = sum((c.width.cm if c.width else 0) for c in t.rows[0].cells)
        if wsum > max(s["band"] for s in secs) + 0.1:
            over.append((i, round(wsum, 1)))
    text = "\n".join([p.text for p in d.paragraphs] +
                     [c.text for t in d.tables for r in t.rows for c in r.cells])
    bad = ["%s x%d" % (lbl, text.count(ch)) for ch, lbl in BAD.items() if text.count(ch)]
    shaded = sum(1 for t in d.tables
                 if re.search(r'w:fill="(?!auto)[0-9A-Fa-f]{6}"', t.rows[0].cells[0]._tc.xml))
    repeat = sum(1 for t in d.tables if t.rows[0]._tr.find(".//" + W + "tblHeader") is not None)
    return {"doc": d, "tables": tables, "rows": rows, "styles": styles, "sections": secs,
            "over": over, "bad": bad, "shaded": shaded, "repeat": repeat}


def run_word(docx, export_pdf=None):
    export = ""
    if export_pdf:
        export = '$doc.ExportAsFixedFormat("%s", 17)' % export_pdf.replace("\\", "\\\\")
    script = PS_UPDATE.format(docx=docx.replace("\\", "\\\\"), export=export)
    exe = "pwsh"
    try:
        out = subprocess.run([exe, "-NoProfile", "-Command", script],
                             capture_output=True, text=True, timeout=600,
                             encoding="utf-8", errors="replace")
    except FileNotFoundError:
        out = subprocess.run(["powershell", "-NoProfile", "-Command", script],
                             capture_output=True, text=True, timeout=600,
                             encoding="utf-8", errors="replace")
    if out.returncode != 0:
        print("  Word недоступен или вернул ошибку:")
        print("  " + (out.stderr or "").strip()[:400])
        return None
    return (out.stdout or "").strip()


def render_png(pdf, pages, outdir):
    try:
        import pymupdf
    except ImportError:
        print("  Нет pymupdf: pip install pymupdf - без него картинок страниц не будет.")
        return []
    os.makedirs(outdir, exist_ok=True)
    d = pymupdf.open(pdf)
    made = []
    for n in pages:
        if 1 <= n <= d.page_count:
            f = os.path.join(outdir, "page_%02d.png" % n)
            d[n - 1].get_pixmap(dpi=110).save(f)
            made.append(f)
    return made


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--md", help="исходный markdown для сверки количества таблиц и строк")
    ap.add_argument("--sample", help="образец, с оформлением которого сравнить")
    ap.add_argument("--update-fields", action="store_true", help="собрать оглавление через Word")
    ap.add_argument("--png", help="номера страниц через запятую: 1,2,9")
    ap.add_argument("--outdir", default="preview", help="куда класть картинки страниц")
    a = ap.parse_args()

    rep = docx_report(a.docx)
    print("=== %s ===" % os.path.basename(a.docx))
    print("  таблиц %d, строк в таблицах %d" % (len(rep["tables"]), rep["rows"]))
    parts = []
    for i, s in enumerate(rep["sections"]):
        note = ""
        if not s["footer"]:
            note = " титульная, без подвала" if (i == 0 and s["first_differs"]) else " БЕЗ ПОДВАЛА"
        parts.append("%s полоса %s см%s" % (s["orient"], s["band"], note))
    print("  секции: %s" % ", ".join(parts))
    print("  шапок с заливкой %d из %d, с повтором при переносе %d"
          % (rep["shaded"], len(rep["tables"]), rep["repeat"]))
    print("  стили: %s" % ", ".join("%s=%d" % kv for kv in sorted(rep["styles"].items())))
    if rep["over"]:
        print("  ТАБЛИЦЫ ШИРЕ ПОЛОСЫ: %s" % rep["over"])
    if rep["bad"]:
        print("  ЗАПРЕЩЕННЫЕ СИМВОЛЫ: %s" % ", ".join(rep["bad"]))

    if a.md:
        t_md, r_md = md_stats(a.md)
        ok = (t_md == len(rep["tables"]) and r_md == rep["rows"])
        print("  сверка с markdown: таблиц %d/%d, строк %d/%d - %s"
              % (t_md, len(rep["tables"]), r_md, rep["rows"], "СОВПАДАЕТ" if ok else "РАСХОЖДЕНИЕ"))
        if not ok:
            print("    Обычная причина: две таблицы подряд без абзаца между ними "
                  "склеиваются Word в одну.")

    if a.sample:
        s = docx_report(a.sample)
        print("  образец: таблиц %d, стили: %s"
              % (len(s["tables"]), ", ".join(sorted(s["styles"]))[:120]))
        only_here = set(rep["styles"]) - set(s["styles"])
        if only_here:
            print("    стили, которых нет в образце: %s" % ", ".join(sorted(only_here)))

    pdf = None
    if a.update_fields or a.png:
        pdf = os.path.splitext(a.docx)[0] + ".preview.pdf" if a.png else None
        out = run_word(os.path.abspath(a.docx), os.path.abspath(pdf) if pdf else None)
        if out:
            print("  Word: %s" % out)

    if a.png and pdf and os.path.exists(pdf):
        pages = [int(x) for x in a.png.split(",") if x.strip().isdigit()]
        made = render_png(pdf, pages, a.outdir)
        for f in made:
            print("  картинка: %s" % f)
        print("  Посмотреть картинки глазами обязательно: верстка проверяется только так.")


if __name__ == "__main__":
    sys.exit(main())
